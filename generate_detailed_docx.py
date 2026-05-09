import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# --- STYLING ---
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# --- TYTUŁ ---
doc.add_paragraph('\n\n\n')
title = doc.add_heading('DOKUMENTACJA TECHNICZNA PROJEKTU', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Platforma Współpracy Sieciowej (GraphQL, WebSocket, Kafka, Raw TCP)', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Author / subject details should be added manually before final delivery if required.

doc.add_page_break()

# --- ROZDZIAŁ 1 ---
doc.add_heading('1. Wprowadzenie i Cel Projektu', level=1)
doc.add_paragraph('Celem projektu była implementacja Platformy Współpracy Sieciowej wykorzystującej zróżnicowane '
                  'protokoły warstwy aplikacji. System demonstruje praktyczne połączenie synchronicznej wymiany danych '
                  '(GraphQL over HTTPS), komunikacji w czasie rzeczywistym (WebSockets), asynchronicznego przesyłania '
                  'wiadomości (Apache Kafka) oraz własnej implementacji opartej o surowe gniazda TCP (Raw Sockets). '
                  'Interfejs aplikacyjny został w całości zrealizowany w technologii GraphQL, co eliminuje problemy '
                  'over-fetchingu i under-fetchingu charakterystyczne dla klasycznego REST.')

# --- ROZDZIAŁ 2 ---
doc.add_heading('2. Architektura Systemu', level=1)
doc.add_paragraph('System został zaprojektowany z myślą o architekturze zorientowanej na usługi (SOA) i skonteneryzowany '
                  'przy użyciu narzędzia Docker (plik docker-compose.yml). Składa się on z następujących głównych komponentów:')
doc.add_paragraph('- API Gateway / Główne Serwery (FastAPI + Strawberry GraphQL): Obsługuje zapytania GraphQL, mutacje oraz połączenia WebSockets.', style='List Bullet')
doc.add_paragraph('- Broker Wiadomości (Apache Kafka + Zookeeper): Odpowiada za rozgłaszanie zdarzeń systemowych.', style='List Bullet')
doc.add_paragraph('- Mikroserwisy Konsumujące: Skrypty w tle (analytics_consumer.py, notification_consumer.py) przetwarzające strumienie danych.', style='List Bullet')
doc.add_paragraph('- Serwer TCP (Custom Protocol): Osobny proces nasłuchujący na porcie 9000 obsługujący własny protokół komunikacyjny.', style='List Bullet')
doc.add_paragraph('- Frontend (app/ui/index.html): Minimalistyczny interfejs użytkownika komunikujący się wyłącznie przez GraphQL.', style='List Bullet')

doc.add_paragraph('\nDiagram Architektury (Uproszczony przepływ logiki):')
diagram = """
+------------------+         GraphQL / HTTPS       +---------------------+
|                  | ---------------------------> |                     |
|  Klient (Baza /  |         WebSockets           |    API (FastAPI)    | ---> [SQLite]
|  Przeglądarka)   | <--------------------------- |  (Port: 8443 / TLS) |
|                  |                              +---------------------+
+------------------+                                        |  Producent
        |                                                   v
        |                                         +---------------------+
        | Raw TCP (Port 9000)                     |    Apache Kafka     | <--> Zookeeper
        v                                         |  (tasks/security)   |
+------------------+                              +---------------------+
|   Serwer TCP     |                                        |  Konsument
| (Autorski Prot.) |                                        v
+------------------+                              +---------------------+
                                                  |   Consumer Workers  |
                                                  | (Python background) |
                                                  +---------------------+
"""
doc.add_paragraph(diagram)

# --- ROZDZIAŁ 3 ---
doc.add_heading('3. Szczegóły Komponentów Komunikacyjnych', level=1)

doc.add_heading('3.1 API: GraphQL', level=2)
doc.add_paragraph('Główny backend został napisany w frameworku FastAPI z wykorzystaniem biblioteki Strawberry GraphQL. '
                  'Wszystkie operacje aplikacyjne (rejestracja, logowanie, zarządzanie zadaniami, audyt) realizowane są '
                  'przez jednolity endpoint /graphql z metodami POST. Uwierzytelnienie oparte jest o token sesyjny '
                  'przekazywany w nagłówku X-Auth-Token.')
doc.add_paragraph('Dostępne operacje GraphQL:')
p = doc.add_paragraph()
p.add_run('Query: me, tasks, auditLogs (wymaga roli admin)\n')
p.add_run('Mutation: register, login, createTask')

doc.add_heading('3.2 Komunikacja w Czasie Rzeczywistym: WebSockets', level=2)
doc.add_paragraph('Aby informować klienta o natychmiastowych zmianach w systemie (np. utworzeniu zadania) wdrożony został '
                  'protokół WebSockets (WSS over TLS). Utrzymuje on ciągłość połączenia (full-duplex) unikając narzutu '
                  'nawiązywania połączenia HTTP przy każdym żądaniu.')

doc.add_heading('3.3 Autorski Protokół TCP', level=2)
doc.add_paragraph('Przygotowano autorski serwer TCP uruchamiany w architekturze wielowątkowej. '
                  'Nasłuchuje on na zdefiniowanym w konfiguracji porcie 9000.')
doc.add_paragraph('Specyfikacja Protokołu Komunikacyjnego:')
p = doc.add_paragraph()
p.add_run('1. NICK:<nazwa> – inicjuje sesję użytkownika.\n')
p.add_run('2. MSG:<treść> – wysyła wiadomość do wszystkich podłączonych klientów.\n')
p.add_run('3. QUIT – zamyka gniazdo i kończy połączenie.')

doc.add_heading('3.4 Apache Kafka', level=2)
doc.add_paragraph('Architektura systemu używa kolejki zdarzeń Apache Kafka wspieranej przez Zookeepera. '
                  'Usługi (GraphQL API, TCP Server) pełnią pozycję tzw. Producentów.')
doc.add_paragraph('Główne tematy (Topics):')
doc.add_paragraph('- tasks.events – odpowiedzialny za przesyłanie zmian stanu zasobów biznesowych (dodanie zadania).', style='List Bullet')
doc.add_paragraph('- security.audit – odpowiedzialny za audyt systemowy, rejestrowanie nieautoryzowanych prób dostępu itp.', style='List Bullet')
doc.add_paragraph('Zaimplementowano również dwa skrypty "workerów" (analytics_consumer.py, notification_consumer.py), '
                  'które podłączają się do strumieni Kafki i asynchronicznie przetwarzają dane, odciążając serwer HTTP.')

doc.add_heading('4. Konteneryzacja i Uruchomienie', level=1)
doc.add_paragraph('Cały stos uruchamiany jest za pomocą Docker Compose. Każdy komponent (API, TCP, Kafka, Zookeeper, '
                  'konsumenci) działa w osobnym kontenerze. Zastosowano healthchecki, restart policies oraz dedykowaną '
                  'sieć mostkową zapewniającą izolację i stabilność komunikacji między usługami.')
p = doc.add_paragraph()
p.add_run('docker compose up --build    # budowa i start\n')
p.add_run('docker compose down -v       # zatrzymanie i wyczyszczenie danych\n')
p.add_run('docker compose run --rm -e PYTHONPATH=/app api pytest -q   # testy')

doc.add_heading('5. Testy Automatyczne', level=1)
doc.add_paragraph('Projekt zawiera zestaw testów automatycznych opartych na pytest i TestClient z FastAPI. '
                  'Testy weryfikują poprawność mutacji GraphQL (register, login, createTask), zabezpieczenia RBAC '
                  '(dostęp do audytu tylko dla admina) oraz wymagania autoryzacji WebSocket.')

doc.add_heading('6. Zabezpieczenia (Security)', level=1)
doc.add_paragraph('Wszystkie warstwy stykowe (API) pracują za zasłoną kryptograficzną dostarczaną przez TLS '
                  '(Transport Layer Security). W projekcie załączony został skrypt (scripts/generate_certs.py) '
                  'służący do generowania i odświeżania lokalnej pary certyfikatów klucza publicznego oraz prywatnego. '
                  'Dodatkowo zaimplementowano mechanizm RBAC (Role-Based Access Control) z rolami admin i user oraz '
                  'trwały audit log rejestrujący kluczowe zdarzenia bezpieczeństwa i operacje biznesowe.')

doc.add_heading('7. Podsumowanie', level=1)
doc.add_paragraph('System implementuje pełen stos komunikacji sieciowej od warstwy aplikacji po gniazda TCP. '
                  'Zastosowanie GraphQL jako jedynego interfejsu aplikacyjnego zapewnia elastyczność i spójność komunikacji, '
                  'a konteneryzacja z healthcheckami gwarantuje powtarzalność środowiska niezbędną podczas demonstracji.')

doc.save('Technical_Documentation.docx')
print("Dokument wygenerowany poprawnie.")
