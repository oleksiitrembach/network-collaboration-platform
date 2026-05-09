from __future__ import annotations

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


def set_paragraph_spacing(paragraph, space_after=Pt(6), line_spacing=1.15):
    paragraph.paragraph_format.space_after = space_after
    paragraph.paragraph_format.line_spacing = line_spacing


def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    set_paragraph_spacing(heading, space_after=Pt(8))
    return heading


def add_paragraph_custom(doc, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    p.alignment = alignment
    set_paragraph_spacing(p)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    set_paragraph_spacing(p)
    return p


def main():
    doc = Document()

    # --- STYLE ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # --- TITLE PAGE ---
    doc.add_paragraph("\n\n\n")
    title = doc.add_heading("DOKUMENTACJA TECHNICZNA PROJEKTU", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Arial"
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    subtitle = doc.add_heading("Platforma Współpracy Sieciowej (GraphQL, WebSocket, Kafka, Raw TCP)", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_page_break()

    # --- 1. WPROWADZENIE ---
    add_heading_custom(doc, "1. Wprowadzenie i cel projektu", level=1)
    add_paragraph_custom(doc,
        "Celem projektu była implementacja Platformy Współpracy Sieciowej wykorzystującej zróżnicowane "
        "protokoły warstwy aplikacji. System demonstruje praktyczne połączenie synchronicznej wymiany danych "
        "(GraphQL over HTTPS), komunikacji w czasie rzeczywistym (WebSockets), asynchronicznego przesyłania "
        "wiadomości (Apache Kafka) oraz własnej implementacji opartej o surowe gniazda TCP (Raw Sockets). "
        "Interfejs aplikacyjny został w całości zrealizowany w technologii GraphQL, co eliminuje problemy "
        "over-fetchingu i under-fetchingu charakterystyczne dla klasycznego REST."
    )

    # --- 2. ARCHITEKTURA ---
    add_heading_custom(doc, "2. Architektura systemu", level=1)
    add_paragraph_custom(doc,
        "System został zaprojektowany z myślą o architekturze zorientowanej na usługi (SOA) i skonteneryzowany "
        "przy użyciu narzędzia Docker (plik docker-compose.yml). Składa się on z następujących głównych komponentów:"
    )
    add_bullet(doc, "API Gateway / Główne Serwery (FastAPI + Strawberry GraphQL): Obsługuje zapytania GraphQL, mutacje oraz połączenia WebSockets.")
    add_bullet(doc, "Broker Wiadomości (Apache Kafka + Zookeeper): Odpowiada za rozgłaszanie zdarzeń systemowych.")
    add_bullet(doc, "Mikroserwisy Konsumujące: Skrypty w tle (analytics_consumer.py, notification_consumer.py) przetwarzające strumienie danych.")
    add_bullet(doc, "Serwer TCP (Custom Protocol): Osobny proces nasłuchujący na porcie 9000 obsługujący własny protokół komunikacyjny.")
    add_bullet(doc, "Frontend (app/ui/index.html): Minimalistyczny, monochromatyczny interfejs użytkownika komunikujący się wyłącznie przez GraphQL.")

    add_paragraph_custom(doc, "Diagram architektury (uproszczony przepływ logiki):")
    diagram = """
+------------------+         GraphQL / HTTPS      +---------------------+
|                  | ---------------------------> |                     |
|  Klient (Baza /  |         WebSockets           |    API (FastAPI)    | ---> [SQLite]
|  Przeglądarka)   | <--------------------------- |  (Port: 8443 / TLS) |
|                  |                              +---------------------+
+------------------+                                        |  Producent
        |                                                   v
        |                                         +---------------------+
        | Raw TCP (Port 9000)                     |    Apache Kafka     | <--> Zookeeper
        v                                         |  (tasks/security)   |
+------------------+                              +---------------------+
|   Serwer TCP     |                                        |  Konsument
| (Autorski Prot.) |                                        v
+------------------+                              +---------------------+
                                                  |   Consumer Workers  |
                                                  | (Python background) |
                                                  +---------------------+
    """
    p = doc.add_paragraph()
    run = p.add_run(diagram)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    set_paragraph_spacing(p)

    # --- 3. KOMPONENTY ---
    add_heading_custom(doc, "3. Szczegóły komponentów komunikacyjnych", level=1)

    add_heading_custom(doc, "3.1 API: GraphQL", level=2)
    add_paragraph_custom(doc,
        "Główny backend został napisany w frameworku FastAPI z wykorzystaniem biblioteki Strawberry GraphQL. "
        "Wszystkie operacje aplikacyjne (rejestracja, logowanie, zarządzanie zadaniami, audyt) realizowane są "
        "poprzez jednolity endpoint /graphql z metodami POST. Uwierzytelnienie oparte jest o token sesyjny "
        "przekazywany w nagłówku X-Auth-Token."
    )
    add_paragraph_custom(doc, "Dostępne operacje GraphQL:")
    add_bullet(doc, "Query: me, tasks, auditLogs (wymaga roli admin)")
    add_bullet(doc, "Mutation: register, login, createTask")

    add_heading_custom(doc, "3.2 Komunikacja w czasie rzeczywistym: WebSockets", level=2)
    add_paragraph_custom(doc,
        "Aby informować klienta o natychmiastowych zmianach w systemie (np. utworzeniu zadania) wdrożony został "
        "protokół WebSockets (WSS over TLS). Utrzymuje on ciągłość połączenia (full-duplex) unikając narzutu "
        "nawiązywania połączenia HTTP przy każdym żądaniu."
    )

    add_heading_custom(doc, "3.3 Autorski protokół TCP", level=2)
    add_paragraph_custom(doc,
        "Przygotowano autorski serwer TCP uruchamiany w architekturze wielowątkowej. "
        "Nasłuchuje on na zdefiniowanym w konfiguracji porcie 9000."
    )
    add_paragraph_custom(doc, "Specyfikacja protokołu komunikacyjnego:")
    p = doc.add_paragraph()
    run = p.add_run("1. NICK:<nazwa> – inicjuje sesję użytkownika.\n")
    run.font.name = "Arial"; run.font.size = Pt(11)
    run = p.add_run("2. MSG:<treść> – wysyła wiadomość do wszystkich podłączonych klientów.\n")
    run.font.name = "Arial"; run.font.size = Pt(11)
    run = p.add_run("3. QUIT – zamyka gniazdo i kończy połączenie.")
    run.font.name = "Arial"; run.font.size = Pt(11)
    set_paragraph_spacing(p)

    add_heading_custom(doc, "3.4 Apache Kafka", level=2)
    add_paragraph_custom(doc,
        "Architektura systemu używa kolejki zdarzeń Apache Kafka wspieranej przez Zookeepera. "
        "Usługi (GraphQL API, TCP Server) pełnią pozycję tzw. Producentów."
    )
    add_paragraph_custom(doc, "Główne tematy (topics):")
    add_bullet(doc, "tasks.events – odpowiedzialny za przesyłanie zmian stanu zasobów biznesowych (dodanie zadania).")
    add_bullet(doc, "security.audit – odpowiedzialny za audyt systemowy, rejestrowanie nieautoryzowanych prób dostępu itp.")
    add_paragraph_custom(doc,
        "Zaimplementowano również dwa skrypty workerów (analytics_consumer.py, notification_consumer.py), "
        "które podłączają się do strumieni Kafki i asynchronicznie przetwarzają dane, odciążając serwer HTTP."
    )

    add_heading_custom(doc, "3.5 Bezpieczeństwo", level=2)
    add_paragraph_custom(doc,
        "Wszystkie warstwy stykowe (API) pracują za zasłoną kryptograficzną dostarczaną przez TLS "
        "(Transport Layer Security). W projekcie załączony został skrypt (scripts/generate_certs.py) "
        "służący do generowania i odświeżania lokalnej pary certyfikatów klucza publicznego oraz prywatnego."
    )
    add_paragraph_custom(doc,
        "Dodatkowo zaimplementowano mechanizm RBAC (Role-Based Access Control) z rolami admin i user oraz "
        "trwały audit log rejestrujący kluczowe zdarzenia bezpieczeństwa i operacje biznesowe."
    )

    # --- 4. KONTENERYZACJA ---
    add_heading_custom(doc, "4. Konteneryzacja i uruchomienie", level=1)
    add_paragraph_custom(doc,
        "Cały stos uruchamiany jest za pomocą Docker Compose. Każdy komponent (API, TCP, Kafka, Zookeeper, "
        "konsumenci) działa w osobnym kontenerze. Zastosowano healthchecki, restart policies oraz dedykowaną "
        "sieć mostkową (pas-network) zapewniającą izolację i stabilność komunikacji między usługami."
    )
    add_paragraph_custom(doc, "Podstawowe komendy:")
    p = doc.add_paragraph()
    run = p.add_run("docker compose up --build    # budowa i start\n")
    run.font.name = "Courier New"; run.font.size = Pt(10)
    run = p.add_run("docker compose down -v       # zatrzymanie i wyczyszczenie danych\n")
    run.font.name = "Courier New"; run.font.size = Pt(10)
    run = p.add_run("docker compose run --rm -e PYTHONPATH=/app api pytest -q   # testy")
    run.font.name = "Courier New"; run.font.size = Pt(10)
    set_paragraph_spacing(p)

    # --- 5. TESTY ---
    add_heading_custom(doc, "5. Testy automatyczne", level=1)
    add_paragraph_custom(doc,
        "Projekt zawiera zestaw testów automatycznych opartych na pytest i TestClient z FastAPI. "
        "Testy weryfikują poprawność mutacji GraphQL (register, login, createTask), zabezpieczenia RBAC "
        "(dostęp do audytu tylko dla admina) oraz wymagania autoryzacji WebSocket."
    )

    # --- 6. PODSUMOWANIE ---
    add_heading_custom(doc, "6. Podsumowanie", level=1)
    add_paragraph_custom(doc,
        "System implementuje pełen stos komunikacji sieciowej od warstwy aplikacji po gniazda TCP. "
        "Zastosowanie GraphQL jako jedynego interfejsu aplikacyjnego zapewnia elastyczność i spójność komunikacji, "
        "a konteneryzacja z healthcheckami gwarantuje powtarzalność środowiska niezbędną podczas demonstracji. "
    )

    doc.save("Technical_Documentation.docx")
    print("Dokument wygenerowany poprawnie: Technical_Documentation.docx")


if __name__ == "__main__":
    main()
